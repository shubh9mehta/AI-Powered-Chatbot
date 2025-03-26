"""import json
import re
from docx import Document

# Load the .docx file
doc_path = "Spring 2025 Webinar Transcripts.docx"
document = Document(doc_path)

# Initialize variables
sessions = []
current_session = None
current_speakers = set()

# Regex pattern to match speaker lines (e.g., John: ... or Dr. Smith: ...)
speaker_pattern = re.compile(r"^([A-Z][A-Za-z .'-]+):\s*(.+)")

# Process each paragraph
for para in document.paragraphs:
    text = para.text.strip()

    # Skip empty lines
    if not text:
        continue

    # Identify session titles (all caps with "SESSION" or similar keyword)
    if re.match(r"^[A-Z0-9 :\-]{5,}$", text) and "SESSION" in text:
        # Save previous session if it exists
        if current_session:
            current_session["speakers"] = list(current_speakers)
            sessions.append(current_session)

        # Start a new session
        current_session = {
            "session_title": text.strip(),
            "speakers": [],
            "transcript": []
        }
        current_speakers = set()

    # Identify and extract speaker and text
    elif match := speaker_pattern.match(text):
        speaker = match.group(1).strip()
        speech = match.group(2).strip()

        current_session["transcript"].append({
            "speaker": speaker,
            "text": speech
        })
        current_speakers.add(speaker)

    # Append to last speaker’s text if it's a continuation line
    elif current_session and current_session["transcript"]:
        current_session["transcript"][-1]["text"] += " " + text.strip()

# Save the final session
if current_session:
    current_session["speakers"] = list(current_speakers)
    sessions.append(current_session)

# Save to JSON
output_file = "structured_webinar_transcript.json"
with open(output_file, "w", encoding="utf-8") as f:
    json.dump(sessions, f, indent=2, ensure_ascii=False)

print(f"✅ Transcript converted and saved to: {output_file}")
"""

import json
import re
from docx import Document
import hashlib
from datetime import datetime
from pathlib import Path

# Load the .docx file
doc_path = "Spring 2025 Webinar Transcripts.docx"
output_path = Path("combined_knowledge_base.json")

# Function to process transcript as a single session
def process_webinar_transcript(doc_path):
    try:
        document = Document(doc_path)
    except Exception as e:
        print(f"❌ Error opening document {doc_path}: {e}")
        return []

    # Initialize a single session
    session = {
        "session_title": "Spring 2025 Research Data Webinar",
        "session_date": "Spring 2025",
        "speakers": [],
        "transcript": [],
        "source_file": doc_path,
        "processed_at": datetime.now().isoformat()
    }

    speakers_set = set()

    # Pattern to detect speaker lines
    speaker_pattern = re.compile(r"^([A-Z][A-Za-z .'-]+):\s*(.+)")

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Identify speaker lines
        if match := speaker_pattern.match(text):
            speaker = match.group(1).strip()
            speech = match.group(2).strip()

            session["transcript"].append({
                "speaker": speaker,
                "text": speech,
                "timestamp": None
            })
            speakers_set.add(speaker)

        # Handle continuation lines
        elif session["transcript"]:
            session["transcript"][-1]["text"] += " " + text

    # Finalize session details
    session["speakers"] = list(speakers_set)
    session["word_count"] = sum(len(item["text"].split()) for item in session["transcript"])
    session["_id"] = hashlib.md5(json.dumps(session, sort_keys=True).encode()).hexdigest()

    return [session] if session["transcript"] else []

# Process the transcript
sessions = process_webinar_transcript(doc_path)

# Load or initialize the combined knowledge base
if output_path.exists():
    with open(output_path, "r", encoding="utf-8") as f:
        combined_knowledge_base = json.load(f)
else:
    combined_knowledge_base = {}

# Replace the "Webinar_Transcripts" section with the new data
if sessions:
    combined_knowledge_base["Webinar_Transcripts"] = sessions

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(combined_knowledge_base, f, indent=2, ensure_ascii=False)

    print(f"✅ Webinar transcript successfully added to: {output_path}")
    print(f"📦 Sessions processed: {len(sessions)} | Total words: {sum(s['word_count'] for s in sessions)}")
else:
    print("⚠️ No valid webinar data found in the document.")
